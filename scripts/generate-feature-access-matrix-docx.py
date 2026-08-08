from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "FEATURE_ACCESS_MATRIX_READABLE.docx"

INK = "102033"
NAVY = "0B2545"
TEAL = "0F6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_TEAL = "EAF6F5"
LIGHT_GRAY = "F2F4F7"
MUTED = "536270"
WHITE = "FFFFFF"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = props.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        props.append(shade)
    shade.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D6DEE7", size="6"):
    props = cell._tc.get_or_add_tcPr()
    borders = props.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_fixed_table_layout(table, widths):
    table.autofit = False
    table_props = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_props.append(layout)
    table_width = sum(widths)
    tbl_w = table_props.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_props.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def font_run(run, size=9.3, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(paragraph, before=0, after=0, line=1.05, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(paragraph, text, size=9.3, bold=False, color=INK):
    run = paragraph.add_run(text)
    font_run(run, size=size, bold=bold, color=color)
    return run


def fill_cell(cell, text, *, header=False, center=False, fill=None, size=8.4, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)
    set_cell_margins(cell)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT)
    add_text(p, text, size=size, bold=header or bold, color=WHITE if header else INK)


def add_table(doc, headers, rows, widths, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell, text in zip(table.rows[0].cells, headers):
        fill_cell(cell, text, header=True, center=True, fill=NAVY, size=8.1)
    set_repeat_table_header(table.rows[0])
    for index, row_values in enumerate(rows):
        row = table.add_row()
        for column, (cell, text) in enumerate(zip(row.cells, row_values)):
            fill = LIGHT_GRAY if index % 2 else WHITE
            center = column >= len(row_values) - 3 and len(row_values) >= 4
            fill_cell(cell, text, center=center, fill=fill, size=7.9 if compact else 8.25)
    set_fixed_table_layout(table, widths)
    p = doc.add_paragraph()
    set_para_format(p, after=4)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para_format(p, before=13 if level == 1 else 9, after=5)
    r = p.add_run(text)
    font_run(r, size=15 if level == 1 else 11.5, bold=True, color=TEAL if level == 1 else NAVY)
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, text, *, color=INK, size=9.7, after=6, bold_prefix=None):
    p = doc.add_paragraph()
    set_para_format(p, after=after, line=1.15)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=size, bold=True, color=color)
        add_text(p, text[len(bold_prefix):], size=size, color=color)
    else:
        add_text(p, text, size=size, color=color)
    return p


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_fixed_table_layout(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_border(cell, color="91C8C5", size="8")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    set_para_format(p, line=1.15)
    add_text(p, f"{label}  ", size=9.7, bold=True, color=TEAL)
    add_text(p, text, size=9.7, color=INK)
    spacer = doc.add_paragraph()
    set_para_format(spacer, after=3)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(p)
    add_text(p, "TrAIner · Matriz de Acesso · 07 ago 2026", size=8, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    add_footer(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    set_para_format(title, after=3)
    add_text(title, "Matriz de Licenças e Funcionalidades", size=23, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    set_para_format(subtitle, after=10)
    add_text(subtitle, "TrAIner · referência comercial e operacional de acesso", size=11.5, color=MUTED)

    meta = doc.add_table(rows=1, cols=3)
    set_fixed_table_layout(meta, [1900, 3900, 3560])
    for cell, label, value in zip(meta.rows[0].cells, ("VERSÃO", "ESTADO", "FONTE TÉCNICA"), ("1.3 · 07/08/2026", "Referência de acesso efectiva", "docs/FEATURE_ACCESS_MATRIX.md")):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell, color="CBD5E1")
        set_cell_margins(cell, top=100, bottom=100)
        p = cell.paragraphs[0]
        set_para_format(p)
        add_text(p, label + "\n", size=7.4, bold=True, color=MUTED)
        add_text(p, value, size=9.1, bold=True, color=NAVY)

    callout(doc, "LEITURA RÁPIDA", "FREE tem 1 geração autónoma por semana e até 6 exercícios. AI FITNESS e AI PERFORMANCE têm geração autónoma ilimitada, sujeita ao Uso Justo. Categoria de exercício e dias do plano prescrito não são gates comerciais. AI PERFORMANCE diferencia-se por análise avançada e métricas de desempenho.")

    heading(doc, "1. Planos de aluno", 1)
    client_rows = [
        ("Geração de treino por IA", "A IA cria treino autónomo", "Sim", "Sim", "Sim"),
        ("Sessões por semana", "Gerações autónomas", "1", "Ilimitado*", "Ilimitado*"),
        ("Exercícios por sessão", "Teto comercial por sessão IA", "6", "Sem limite", "Sem limite"),
        ("Categoria do treino", "Objectivo, perfil e regras de segurança", "Sem gate comercial", "Sem gate comercial", "Sem gate comercial"),
        ("Plano TRAINER vinculado", "Dias e exercícios do plano prescrito", "Integral", "Integral", "Integral"),
        ("Check-in rápido", "Registo de prontidão simples", "Sim", "Sim", "Sim"),
        ("Check-in completo", "Dados detalhados, texto/voz", "Não", "Sim", "Sim"),
        ("Métricas fitness avançadas", "Força, volume e tendências", "Bloqueadas", "Sim", "Sim"),
        ("Métricas de desempenho", "ATL/CTL/TSB e indicadores atléticos", "Bloqueadas", "Bloqueadas", "Sim"),
        ("AI Check-in Adjustment", "Ajuste ao estado diário", "Não", "Sim", "Sim"),
        ("AI Advanced Analysis", "Análise preditiva de carga/recuperação", "Não", "Não", "Sim"),
    ]
    add_table(doc, ["Funcionalidade", "O que faz", "FREE", "AI FITNESS", "AI PERFORMANCE"], client_rows, [1800, 3400, 1200, 1480, 1480])
    body(doc, "* Uso ilimitado para utilização pessoal normal, sujeito à Política de Uso Justo e aos Termos de Uso. Não há contador comercial visível.", size=8.3, color=MUTED, after=3)

    heading(doc, "2. Regras de acesso que evitam promessas incorretas", 1)
    body(doc, "Geração autónoma: os limites FREE aplicam-se apenas ao treino que o próprio aluno gera com IA. AI FITNESS e AI PERFORMANCE não são modelos de IA distintos; partilham o mesmo pipeline, com entitlements e contexto analítico diferentes.", bold_prefix="Geração autónoma:")
    body(doc, "Plano prescrito por TRAINER: um aluno com vínculo activo executa todos os dias e exercícios do plano que recebeu. A licença do aluno não reduz nem filtra esse plano. Continuam aplicáveis regras de segurança e adequação clínica/física.", bold_prefix="Plano prescrito por TRAINER:")
    body(doc, "Diferencial AI PERFORMANCE: análise avançada, métricas de desempenho e contexto preditivo. Não é uma permissão exclusiva para uma categoria de exercícios.", bold_prefix="Diferencial AI PERFORMANCE:")

    heading(doc, "3. Planos de TRAINER", 1)
    trainer_rows = [
        ("Clientes activos", "Alunos geridos simultaneamente", "3", "5 / 15 / 30", "Ilimitado*"),
        ("Coach DNA", "Metodologia do treinador aplicada à IA", "Não", "Sim", "Sim"),
        ("AI Score avançado", "Visão completa por aluno", "Não", "Sim", "Sim"),
        ("AI Check-in Adjustment", "Ajuste de planos por check-in", "Não", "Sim", "Sim"),
        ("AI Advanced Analysis", "Análises preditivas dos alunos", "Não", "Sim", "Sim"),
        ("Studio Branding", "Marca própria no espaço do aluno", "Não", "Em breve", "Em breve"),
        ("Marketplace", "Listagem e revenue share", "Não", "Não", "Em breve"),
        ("Dashboard do cliente", "Visualização de dados do aluno", "Override completo", "Override completo", "Override completo"),
    ]
    add_table(doc, ["Funcionalidade", "O que faz", "TRIAL", "PRO", "ELITE"], trainer_rows, [1800, 3400, 1200, 1480, 1480])
    body(doc, "* Uso ilimitado sujeito à Política de Uso Justo.", size=8.3, color=MUTED, after=4)

    doc.add_page_break()
    heading(doc, "4. Direitos patrocinados pelo TRAINER", 1)
    body(doc, "A licença do TRAINER patrocina ao aluno FREE apenas capacidades determinísticas e sem custo de inferência. Ela não transfere voz, interpretação ou ajuste por IA.")
    sponsored_rows = [
        ("Check-in manual detalhado", "Sim", "Captura estruturada completa além do check-in rápido."),
        ("Dados operacionais para o TRAINER", "Sim", "Aderência, frequência, carga e volume."),
        ("Check-in por voz", "Não", "Requer `checkin.voice_input` da licença do aluno."),
        ("Interpretação de check-in por IA", "Não", "Requer entitlement de IA do próprio aluno."),
        ("Ajuste de plano por IA", "Não", "Requer `ai.checkin_adjustment` da licença do aluno."),
    ]
    add_table(doc, ["Recurso para aluno FREE vinculado", "Disponível", "Regra"], sponsored_rows, [3000, 1300, 5060])

    heading(doc, "5. Autoridade técnica e evidência", 1)
    authority_rows = [
        ("`workout.sessions_per_week`", "FREE = 1 · AI FITNESS = ∞ · AI PERFORMANCE = ∞", "Aplicado no servidor"),
        ("`workout.exercises_per_session`", "FREE = 6 · restantes = ∞", "Aplicado no servidor e fallback"),
        ("`workout.exercise_type`", "Legado; já não é gate comercial", "Não lido; `fitnessOnly = false`"),
        ("`trainer_plan.days_per_week`", "Legado; plano prescrito não é reduzido por tier", "Não lido"),
        ("`ai.checkin_adjustment`", "FREE = não · AI FITNESS/PERFORMANCE = sim", "Aplicado"),
        ("`ai.advanced_analysis`", "Somente AI PERFORMANCE", "Aplicado"),
    ]
    add_table(doc, ["Feature key", "Configuração efectiva", "Estado"], authority_rows, [2500, 4300, 2560])

    heading(doc, "6. Nota de manutenção", 1)
    body(doc, "Este documento é uma versão de leitura da matriz técnica. A fonte de verdade versionada permanece `docs/FEATURE_ACCESS_MATRIX.md`. Qualquer alteração de licença, entitlement, patrocínio TRAINER, custo de IA ou promessa pública deve seguir `docs/AI_GOVERNANCE_CHANGE_GATE.md` e actualizar a matriz no mesmo conjunto de mudanças.")
    body(doc, "Históricos de decisões e de testes não descrevem necessariamente o estado actual; são identificados como históricos e remetem para a matriz de referência.", color=MUTED, size=8.8, after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
